from io import BytesIO

from docx import Document

from app.services.parsers import detect_and_parse, parse_csv, parse_docx, parse_txt


def test_parse_txt() -> None:
    assert parse_txt(b"hello") == "hello"


def test_detect_and_parse_txt() -> None:
    assert detect_and_parse("notes.txt", b"line item") == "line item"


def test_parse_docx() -> None:
    document = Document()
    document.add_paragraph("First line")
    document.add_paragraph("Second line")
    buffer = BytesIO()
    document.save(buffer)

    assert parse_docx(buffer.getvalue()) == "First line\nSecond line"


def test_parse_csv() -> None:
    payload = b"name,value\ncoverage,2.8x\n"

    result = parse_csv(payload)

    assert "coverage" in result
    assert "2.8x" in result


def test_detect_and_parse_docx() -> None:
    document = Document()
    document.add_paragraph("Discovery notes")
    buffer = BytesIO()
    document.save(buffer)

    assert detect_and_parse("notes.docx", buffer.getvalue()) == "Discovery notes"
